# =============================================================================
# document_generator.py
# PURPOSE: Build the Analysis Word document matching the EXACT format
#          of the original uploaded document.
#
# EXACT SPECS FROM ORIGINAL:
#   - A4 page (11906 x 16838 DXA), all margins 1 inch (1440 DXA)
#   - Header paragraphs: bold, 12.5pt (177800 EMU = ~12.5pt), centered
#   - Title box: full-width table (11340 DXA), no visible cell borders,
#                bold + italic text, centered
#   - Data table: 7783 DXA wide, centered on page
#                 Col widths: 983 | 1700 | 1700 | 1700 | 1700 (DXA)
#                 Font: 10pt (sz=20), black borders sz=4
#   - Charts: one per paragraph, centered, full width
#             placed one after another below the table
# =============================================================================

import os
import zipfile
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# -----------------------------------------------------------------------
# HEADER & FOOTER
# Both documents share the same header and footer:
#   Header: "Techno Main Salt Lake       EM-4/1, Sector-V, Salt Lake City, Kolkata- 700091"
#           bold italic, 11pt, left aligned
#   Footer: "© TMSL" on left | page number on right
#           bold italic, 11pt
# -----------------------------------------------------------------------

def _inject_header_footer(docx_path: str):
    """
    Inject the exact TMSL header and footer into a saved docx file by direct
    zip manipulation. The header/footer XMLs are stored in core/../assets/ and
    contain complex anchored drawings (horizontal lines + right-aligned text)
    that cannot be recreated via the python-docx API.

    Header: top-right italic two lines — "Techno Main Salt Lake" /
            "EM-4/1, Sector-V, Salt Lake City, Kolkata- 700091"
            with a thick line below

    Footer: thick line above — "© TMSL" italic left | "Page X of 2" italic right
    """
    import re as _re, io as _io, shutil as _shutil, tempfile as _tmp

    assets_dir = os.path.join(os.path.dirname(__file__), '..', 'assets')
    hdr_path = os.path.join(assets_dir, 'header1.xml')
    ftr_path = os.path.join(assets_dir, 'footer1.xml')
    img_path = os.path.join(assets_dir, 'image5.png')

    if not os.path.exists(hdr_path):
        print("  ⚠  assets/header1.xml not found — header/footer skipped")
        return

    with open(hdr_path, 'rb') as f: header_xml = f.read()
    with open(ftr_path, 'rb') as f: footer_xml = f.read()
    with open(img_path, 'rb') as f: image5     = f.read()

    hf_rels = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r'
        b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        b'<Relationship Id="rId1" '
        b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
        b'Target="media/image5.png"/>'
        b'</Relationships>'
    )

    # Read all files from the docx zip
    with zipfile.ZipFile(docx_path, 'r') as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    # Inject header/footer XML and supporting image
    files['word/header1.xml']            = header_xml
    files['word/footer1.xml']            = footer_xml
    files['word/_rels/header1.xml.rels'] = hf_rels
    files['word/_rels/footer1.xml.rels'] = hf_rels
    files['word/media/image5.png']       = image5

    # Add header/footer relationships to document.xml.rels
    doc_rels = files['word/_rels/document.xml.rels'].decode('utf-8')
    existing = _re.findall(r'Id="rId(\d+)"', doc_rels)
    max_id   = max(int(x) for x in existing) if existing else 10
    hdr_rid  = f"rId{max_id + 1}"
    ftr_rid  = f"rId{max_id + 2}"
    img_rid  = f"rId{max_id + 3}"
    new_rels = (
        f'<Relationship Id="{hdr_rid}" '
        f'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/header" '
        f'Target="header1.xml"/>'
        f'<Relationship Id="{ftr_rid}" '
        f'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer" '
        f'Target="footer1.xml"/>'
        f'<Relationship Id="{img_rid}" '
        f'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
        f'Target="media/image5.png"/>'
    )
    doc_rels = doc_rels.replace('</Relationships>', new_rels + '</Relationships>')
    files['word/_rels/document.xml.rels'] = doc_rels.encode('utf-8')

    # Wire header/footer into sectPr in document.xml
    doc_xml = files['word/document.xml'].decode('utf-8')
    doc_xml = _re.sub(r'<w:headerReference[^/]*/>', '', doc_xml)
    doc_xml = _re.sub(r'<w:footerReference[^/]*/>', '', doc_xml)
    hf_refs = (
        f'<w:headerReference r:id="{hdr_rid}" w:type="default"/>'
        f'<w:footerReference r:id="{ftr_rid}" w:type="default"/>'
    )
    doc_xml = _re.sub(r'(<w:sectPr(?:\s[^>]*)?>)', r'\1' + hf_refs, doc_xml, count=1)
    files['word/document.xml'] = doc_xml.encode('utf-8')

    # Write back to the same file via a temp file
    tmp = docx_path + '.tmp'
    with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)
    _shutil.move(tmp, docx_path)


# -----------------------------------------------------------------------
# LOW-LEVEL XML HELPERS
# -----------------------------------------------------------------------

def _set_table_borders(table, val="single", sz="4", color="000000"):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), val)
        border.set(qn("w:sz"), sz)
        border.set(qn("w:color"), color)
        border.set(qn("w:space"), "0")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _set_cell_borders_nil(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "nil")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:color"), "000000")
        border.set(qn("w:space"), "0")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_table_width_centered(table, width_dxa: int):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(width_dxa))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)


def _set_col_width(cell, width_dxa: int):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_cell_valign(cell, val="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), val)
    tcPr.append(vAlign)


def _set_cell_shading(cell, fill="auto"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


# -----------------------------------------------------------------------
# CHART GENERATION
# Colors extracted exactly from the original NBA report document
# -----------------------------------------------------------------------

COLOR_HIGH     = "#2864c8"   # Blue
COLOR_MODERATE = "#f08c00"   # Amber
COLOR_LOW      = "#dc2800"   # Red
COLOR_TEXT     = "#141414"   # Near-black


def generate_pie_chart(co: dict, output_dir: str) -> str:
    """
    Generate one pie chart per CO — same layout as v1 but with exact
    colors from the original document.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    LABELS = ["High", "Moderate", "Low"]
    PCTS   = [co["high_pct"], co["moderate_pct"], co["low_pct"]]
    COLORS = [COLOR_HIGH, COLOR_MODERATE, COLOR_LOW]

    # Only include slices with value > 0
    filtered = [(l, p, c) for l, p, c in zip(LABELS, PCTS, COLORS) if p > 0]
    labels_f = [x[0] for x in filtered]
    pcts_f   = [x[1] for x in filtered]
    colors_f = [x[2] for x in filtered]

    fig, ax = plt.subplots(figsize=(6, 4.5))
    fig.patch.set_facecolor("white")

    wedges, texts, autotexts = ax.pie(
        pcts_f,
        labels=None,
        colors=colors_f,
        autopct="%1.0f%%",
        startangle=90,
        wedgeprops={"edgecolor": "white", "linewidth": 2},
        pctdistance=0.65,
    )

    # White bold % inside slices
    for autotext in autotexts:
        autotext.set_color("white")
        autotext.set_fontweight("bold")
        autotext.set_fontsize(12)

    # Legend below chart
    ax.legend(
        wedges,
        [f"{l} – {p}%" for l, p in zip(labels_f, pcts_f)],
        loc="lower center",
        bbox_to_anchor=(0.5, -0.18),
        ncol=len(labels_f),
        fontsize=10,
        frameon=False,
        labelcolor=COLOR_TEXT,
    )

    # Title: CO code + description
    short_desc = co["description"][:65] + ("..." if len(co["description"]) > 65 else "")
    ax.set_title(
        f"{co['code']}\n{short_desc}",
        fontsize=10,
        fontweight="bold",
        color=COLOR_TEXT,
        pad=12,
    )

    plt.tight_layout()

    safe_code = co["code"].replace("/", "-").replace("\\", "-").replace(":", "-")
    filepath = os.path.join(output_dir, f"chart_{safe_code}.png")
    plt.savefig(filepath, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return filepath


def generate_all_charts(co_data: list, output_dir: str) -> list:
    """Generate one chart per CO. Returns list of PNG paths."""
    os.makedirs(output_dir, exist_ok=True)
    paths = []
    for co in co_data:
        path = generate_pie_chart(co, output_dir)
        paths.append(path)
        print(f"  ✓ Chart: {path}")
    return paths


# -----------------------------------------------------------------------
# ANALYSIS DOCUMENT BUILDER
# -----------------------------------------------------------------------

def build_analysis_document(co_data, chart_paths, course_info, output_path):
    doc = Document()

    # A4 page, 1 inch margins (1440 DXA = 1 inch = 2.54cm)
    section = doc.sections[0]
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

    # Standard TMSL header and footer injected after save (see _inject_header_footer)

    dept        = course_info.get("department", "Department of Computer Science and Engineering (CSE)")
    acad_year   = course_info.get("academic_year", "N/A")
    n_resp      = course_info.get("respondent_count", 0)
    course_code = course_info.get("course_code", "")
    course_name = course_info.get("course_name", "")

    # ---- Helper: add a centered bold header paragraph (exact 12.5pt) ----
    def add_header_para(text, bold=True):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Emu(177800)  # 177800 EMU = 12.5pt (matches original exactly)
        return p

    # ================================================================
    # SECTION 1 — Header paragraphs
    # ================================================================
    add_header_para(dept)
    doc.add_paragraph()
    add_header_para("Analysis of Course-Exit Feedback (End-Term)")
    doc.add_paragraph()
    add_header_para(f"(CAY: {acad_year})")
    add_header_para(f"No. of Respondents: {n_resp}")

    # Spacing before title box
    for _ in range(3):
        doc.add_paragraph()

    # ================================================================
    # SECTION 2 — Title box
    # Full width (11340 DXA), no visible cell borders, bold+italic, centered
    # Line 1: "% of Course-Exit Feedback Analysis (COURSE_CODE"
    # Line 2: "-Course Name )"
    # ================================================================
    title_table = doc.add_table(rows=1, cols=1)
    _set_table_width_centered(title_table, 11340)
    _set_table_borders(title_table, val="single", sz="4")

    tc = title_table.rows[0].cells[0]
    _set_cell_borders_nil(tc)
    _set_cell_shading(tc, fill="auto")
    _set_cell_valign(tc, "bottom")

    tc.text = ""
    p1 = tc.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("% of Course-Exit Feedback Analysis ")
    r1.bold = True
    r2 = p1.add_run(f"({course_code} ")
    r2.bold = True
    r2.italic = True

    p2 = tc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p2.add_run(f"-{course_name} )")
    r3.bold = True
    r3.italic = True

    doc.add_paragraph()  # spacer

    # ================================================================
    # SECTION 3 — Data table
    # Col widths: 983 DXA (% label) | 1700 DXA × n_cos
    # ================================================================
    num_cos     = len(co_data)
    first_col_w = 983
    co_col_w    = 1700   # matches original exactly

    data_table = doc.add_table(rows=4, cols=num_cos + 1)
    _set_table_width_centered(data_table, first_col_w + co_col_w * num_cos)
    _set_table_borders(data_table, val="single", sz="4", color="000000")

    # Set all column widths
    for row in data_table.rows:
        _set_col_width(row.cells[0], first_col_w)
        for i in range(1, num_cos + 1):
            _set_col_width(row.cells[i], co_col_w)

    def set_cell(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = align
        run = para.add_run(text)
        run.font.size = Pt(10)   # 10pt = sz 20 in XML, matches original
        run.bold = bold
        _set_cell_valign(cell, "center")

    # Header row — CO descriptions
    hdr = data_table.rows[0]
    set_cell(hdr.cells[0], "\u00a0%")   # non-breaking space + %

    for i, co in enumerate(co_data):
        set_cell(hdr.cells[i + 1], f"{co['code']} {co['description']}")

    # High / Moderate / Low rows
    for row_idx, (label, pct_key) in enumerate([
        ("High",     "high_pct"),
        ("Moderate", "moderate_pct"),
        ("Low",      "low_pct"),
    ]):
        row = data_table.rows[row_idx + 1]
        set_cell(row.cells[0], label)
        for i, co in enumerate(co_data):
            set_cell(row.cells[i + 1], str(co[pct_key]),
                     align=WD_ALIGN_PARAGRAPH.CENTER)

    # ================================================================
    # SECTION 4 — Pie charts, one per row, centered, full width
    # Each chart on its own line, stacked vertically one below the other.
    # No spacer after the last chart to avoid a blank trailing page.
    # ================================================================
    doc.add_paragraph()   # small spacer after data table

    valid_charts = [p for p in chart_paths if os.path.exists(p)]

    for idx, chart_path in enumerate(valid_charts):
        # One chart per paragraph, centered
        chart_para = doc.add_paragraph()
        chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = chart_para.add_run()
        run.add_picture(chart_path, width=Cm(14))   # 14cm — fits A4 with margins

        # Spacer between charts only — not after the last one (prevents blank page)
        if idx < len(valid_charts) - 1:
            doc.add_paragraph()

    # ================================================================
    # SECTION 5 — Signature block (bottom RIGHT as per original PDF)
    # Line 1: dotted line ……………………………………………
    # Line 2: (Submitted by – Dr. Biman Roy )
    # Both right-aligned
    # ================================================================
    doc.add_paragraph()
    doc.add_paragraph()

    sig1 = doc.add_paragraph()
    sig1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = sig1.add_run("………………………………………………………………")
    r1.font.size = Pt(11)

    sig2 = doc.add_paragraph()
    sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = sig2.add_run("(Submitted by – Dr. Biman Roy )")
    r2.font.size = Pt(11)

    # ================================================================
    # SAVE
    # ================================================================
    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)
    _inject_header_footer(output_path)
    print(f"  ✓ Analysis document saved: {output_path}")
    return output_path


# -----------------------------------------------------------------------
# ATR DOCUMENT BUILDER
# -----------------------------------------------------------------------

def build_atr_document(
    co_data: list,
    action_taken_texts: list,
    course_info: dict,
    date_range: dict,
    output_path: str
) -> str:
    """
    Build the ATR (Action Taken Report) Word document.

    EXACT SPECS FROM ORIGINAL:
      - Same page as Analysis: A4, 1-inch margins
      - Headers: bold, 12.5pt (177800 EMU), centered
      - Intro paragraph: 9.75pt (139700 EMU), justified
                         date range text is bold + italic
      - Table: 10800 DXA wide, centered
               Col widths: 1017 (Sl.No) | 4507 (Feedback) | 5276 (Action Taken)
               Font: 11pt (sz=22), black borders sz=4

    Structure:
      1. Header paragraphs
      2. Intro sentence with bold-italic date range
      3. Table: Sl.No | Feedback (CO code + description) | Action Taken
    """
    doc = Document()

    # A4 page, 1 inch margins — same as Analysis doc
    section = doc.sections[0]
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

    # Standard TMSL header and footer injected after save (see _inject_header_footer)

    dept        = course_info.get("department", "")
    acad_year   = course_info.get("academic_year", "N/A")
    semester    = course_info.get("semester", "Odd Semester")
    date_str    = date_range.get("range_str", "N/A")

    # ================================================================
    # SECTION 1 — Header paragraphs
    # Exact format: bold, 12.5pt (177800 EMU), centered
    # ================================================================

    def add_header(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Emu(177800)

    add_header("INTERNAL QUALITY ASSURANCE COMMITTEE (IQAC)")
    add_header(f"(CAY: {acad_year})")
    add_header("Feedback Analysis and Action Taken Report")
    add_header("(Course-Exit Feedback)")
    add_header(dept)

    doc.add_paragraph()  # blank line after headers (matches original P5)

    # ================================================================
    # SECTION 2 — Intro paragraph
    # Font: 9.75pt (139700 EMU), justified
    # Date range portion is bold + italic
    # ================================================================

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_intro_run(text, bold=False, italic=False):
        run = intro.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Emu(139700)   # 9.75pt — exact match from original

    add_intro_run("The feedback was collected from the students during ")
    add_intro_run(date_str, bold=True, italic=True)           # date bold+italic
    add_intro_run(f" (for {semester}), post completion of evaluation and assessment.")

    doc.add_paragraph()  # spacer before table

    # ================================================================
    # SECTION 3 — Feedback / Action Taken table
    # 10800 DXA wide, col widths: 1017 | 4507 | 5276
    # Header row + one row per CO
    # Font: 11pt (sz=22), black borders
    # ================================================================

    num_cos = len(co_data)
    table = doc.add_table(rows=num_cos + 1, cols=3)
    _set_table_width_centered(table, 10800)
    _set_table_borders(table, val="single", sz="4", color="000000")

    # Col widths: Sl.No very narrow (~1cm = 565 DXA)
    # Remaining 10235 DXA split: Feedback 40% | Action Taken 60%
    COL_WIDTHS = [565, 4094, 6141]   # sums to 10800 DXA
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            _set_col_width(cell, COL_WIDTHS[i])

    def set_atr_cell(cell, text, bold=False,
                     align=WD_ALIGN_PARAGRAPH.LEFT, valign="center"):
        """Set text in ATR table cell with 11pt font."""
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = align
        run = para.add_run(text)
        run.font.size = Pt(11)   # 11pt = sz 22 in XML — matches original exactly
        run.bold = bold
        _set_cell_valign(cell, valign)

    # Header row
    hdr = table.rows[0]
    set_atr_cell(hdr.cells[0], "Sl. No.",      bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_atr_cell(hdr.cells[1], "Feedback",     bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_atr_cell(hdr.cells[2], "Action Taken", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows — one per CO
    for i, (co, action_text) in enumerate(zip(co_data, action_taken_texts)):
        row = table.rows[i + 1]

        # Sl. No. — sequential number starting from 1
        sl_cell = row.cells[0]
        sl_cell.text = ""
        sl_para = sl_cell.paragraphs[0]
        sl_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sl_run = sl_para.add_run(str(i + 1))
        sl_run.font.size = Pt(11)
        _set_cell_valign(sl_cell, "center")

        # Feedback — CO code + description as plain text, 11pt, justified
        fb_cell = row.cells[1]
        fb_cell.text = ""
        fb_para = fb_cell.paragraphs[0]
        fb_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fb_run = fb_para.add_run(f"{co['code']} {co['description']}")
        fb_run.font.size = Pt(11)
        _set_cell_valign(fb_cell, "center")

        # Action Taken — generated sentence, 11pt, justified
        at_cell = row.cells[2]
        at_cell.text = ""
        at_para = at_cell.paragraphs[0]
        at_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        at_run = at_para.add_run(action_text)
        at_run.font.size = Pt(11)
        _set_cell_valign(at_cell, "center")

    # ================================================================
    # SECTION 4 — Signature block (bottom right)
    # "     ………………………………………..
    #  (Ratified by PAC – Prof. Nairanjana Chowdhury,
    #   Prof. Mrinal Kanti Nath & Subhadra Shaw)"
    # ================================================================
    doc.add_paragraph()
    doc.add_paragraph()

    sig1 = doc.add_paragraph()
    sig1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = sig1.add_run("     ………………………………………..     ")
    r1.font.size = Pt(11)

    sig2 = doc.add_paragraph()
    sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = sig2.add_run("(Ratified by PAC – Prof. Nairanjana Chowdhury,  Prof. Mrinal Kanti Nath & Subhadra Shaw)")
    r2.font.size = Pt(11)

    # ================================================================
    # SAVE
    # ================================================================
    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)
    _inject_header_footer(output_path)
    print(f"  ✓ ATR document saved: {output_path}")
    return output_path
