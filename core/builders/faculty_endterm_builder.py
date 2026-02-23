# =============================================================================
# core/builders/faculty_endterm_builder.py
# PURPOSE: Build Analysis + ATR Word documents for Faculty End-Term feedback.
#
# DIFFERENCES FROM COURSE-EXIT BUILDER:
#   - Analysis title: "Analysis of Faculty Feedback (End-Term) on Teaching and Learning"
#   - Title box: "% of Faculty Feedback Analysis (code name)"
#   - Extra line after title box: "Name of Faculty: Dr. Biman Roy"
#   - Data table: 4 rows (SA / Agree / Neutral / Disagree) instead of 3
#   - Question labels: Q1, Q2... (no CO codes)
#   - ATR subtitle: "(Teaching and Learning – End-Term)"
#
# CHANGING THIS FILE does not affect Course-Exit or Faculty Mid-Term.
# =============================================================================

import os
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.shared.header_footer import inject_header_footer


# -----------------------------------------------------------------------
# XML HELPERS (same as course_exit_builder — duplicated intentionally
# so each builder is fully self-contained and changes don't cross)
# -----------------------------------------------------------------------

def _set_table_borders(table, val="single", sz="4", color="000000"):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   val)
        border.set(qn("w:sz"),    sz)
        border.set(qn("w:color"), color)
        border.set(qn("w:space"), "0")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _set_cell_borders_nil(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "nil")
        border.set(qn("w:sz"),    "0")
        border.set(qn("w:color"), "000000")
        border.set(qn("w:space"), "0")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_table_width_centered(table, width_dxa: int):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(width_dxa))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)


def _set_col_width(cell, width_dxa: int):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_cell_valign(cell, val="center"):
    tc     = cell._tc
    tcPr   = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), val)
    tcPr.append(vAlign)


def _set_cell_shading(cell, fill="auto"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"),  "clear")
    tcPr.append(shd)


def _page_setup(doc):
    section = doc.sections[0]
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)


def _add_header_para(doc, text):
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Emu(177800)
    return p


def _save_and_inject(doc, output_path: str, label: str):
    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)
    inject_header_footer(output_path)
    print(f"  ✓ {label} saved: {output_path}")
    return output_path


# -----------------------------------------------------------------------
# PUBLIC BUILDERS
# -----------------------------------------------------------------------

def build_analysis(question_data: list, chart_paths: list,
                   course_info: dict, output_path: str) -> str:
    """
    Build the Faculty End-Term Analysis document.

    Key differences from course_exit_builder:
      - Title line: "Analysis of Faculty Feedback (End-Term) on Teaching and Learning"
      - Title box: "% of Faculty Feedback Analysis (code name)"
      - Extra paragraph: "Name of Faculty: [name]"
      - Data table has 4 rating rows: SA / Agree / Neutral / Disagree
    """
    doc = Document()
    _page_setup(doc)

    dept         = course_info.get("department", "")
    acad_year    = course_info.get("academic_year", "N/A")
    n_resp       = course_info.get("respondent_count", 0)
    course_code  = course_info.get("course_code", "")
    course_name  = course_info.get("course_name", "")
    faculty_name = course_info.get("faculty_name", "")

    # --- Section 1: Header paragraphs ---
    _add_header_para(doc, dept)
    doc.add_paragraph()
    _add_header_para(doc, "Analysis of Faculty Feedback (End-Term) on Teaching and Learning")
    doc.add_paragraph()
    _add_header_para(doc, f"(CAY: {acad_year})")
    _add_header_para(doc, f"No. of Respondents: {n_resp}")

    for _ in range(3):
        doc.add_paragraph()

    # --- Section 2: Title box ---
    title_table = doc.add_table(rows=1, cols=1)
    _set_table_width_centered(title_table, 11340)
    _set_table_borders(title_table, val="single", sz="4")

    tc = title_table.rows[0].cells[0]
    _set_cell_borders_nil(tc)
    _set_cell_shading(tc, fill="auto")
    _set_cell_valign(tc, "bottom")
    tc.text = ""

    p1 = tc.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("% of Faculty Feedback Analysis ")
    r1.bold = True
    r2 = p1.add_run(f"({course_code} ")
    r2.bold   = True
    r2.italic = True

    p2 = tc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p2.add_run(f"{course_name})")
    r3.bold   = True
    r3.italic = True

    doc.add_paragraph()

    # Faculty name line — unique to faculty feedback
    if faculty_name:
        fn_para = doc.add_paragraph()
        fn_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fn_run = fn_para.add_run(f"Name of Faculty: {faculty_name}")
        fn_run.bold = True
        fn_run.font.size = Emu(177800)
        doc.add_paragraph()

    # --- Section 3: Data table (4 rating rows) ---
    num_qs      = len(question_data)
    first_col_w = 983
    q_col_w     = max(1200, 9026 // max(num_qs, 1))  # dynamic width for many questions

    data_table = doc.add_table(rows=5, cols=num_qs + 1)  # 4 rating rows + header
    _set_table_width_centered(data_table, first_col_w + q_col_w * num_qs)
    _set_table_borders(data_table, val="single", sz="4", color="000000")

    for row in data_table.rows:
        _set_col_width(row.cells[0], first_col_w)
        for i in range(1, num_qs + 1):
            _set_col_width(row.cells[i], q_col_w)

    def set_cell(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = align
        run = para.add_run(text)
        run.font.size = Pt(10)
        run.bold = bold
        _set_cell_valign(cell, "center")

    # Header row: question descriptions
    hdr = data_table.rows[0]
    set_cell(hdr.cells[0], "\u00a0%")
    for i, q in enumerate(question_data):
        set_cell(hdr.cells[i + 1], q["description"])

    # Rating rows
    for row_idx, (label, pct_key) in enumerate([
        ("Strongly Agree", "strongly_agree_pct"),
        ("Agree",          "agree_pct"),
        ("Neutral",        "neutral_pct"),
        ("Disagree",       "disagree_pct"),
    ]):
        row = data_table.rows[row_idx + 1]
        set_cell(row.cells[0], label)
        for i, q in enumerate(question_data):
            set_cell(row.cells[i + 1], str(q[pct_key]),
                     align=WD_ALIGN_PARAGRAPH.CENTER)

    # --- Section 4: Charts ---
    doc.add_paragraph()
    valid = [p for p in chart_paths if os.path.exists(p)]
    for idx, chart_path in enumerate(valid):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(chart_path, width=Cm(14))
        if idx < len(valid) - 1:
            doc.add_paragraph()

    # --- Section 5: Signature ---
    doc.add_paragraph()
    doc.add_paragraph()
    sig1 = doc.add_paragraph()
    sig1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig1.add_run("………………………………………………………………").font.size = Pt(11)
    sig2 = doc.add_paragraph()
    sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig2.add_run("(Submitted by \u2013 Dr. Biman Roy )").font.size = Pt(11)

    return _save_and_inject(doc, output_path, "Analysis document")


def build_atr(question_data: list, action_taken_texts: list,
              course_info: dict, date_range: dict, output_path: str) -> str:
    """
    Build the Faculty End-Term ATR document.

    Key difference from course_exit ATR:
      - Subtitle: "(Teaching and Learning \u2013 End-Term)"
    """
    doc = Document()
    _page_setup(doc)

    dept      = course_info.get("department", "")
    acad_year = course_info.get("academic_year", "N/A")
    semester  = course_info.get("semester", "Odd Semester")
    date_str  = date_range.get("range_str", "N/A")

    # --- Section 1: Headers ---
    _add_header_para(doc, "INTERNAL QUALITY ASSURANCE COMMITTEE (IQAC)")
    _add_header_para(doc, f"(CAY: {acad_year})")
    _add_header_para(doc, "Feedback Analysis and Action Taken Report")
    _add_header_para(doc, "(Teaching and Learning \u2013 End-Term)")
    _add_header_para(doc, dept)
    doc.add_paragraph()

    # --- Section 2: Intro ---
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_run(text, bold=False, italic=False):
        r = intro.add_run(text)
        r.bold   = bold
        r.italic = italic
        r.font.size = Emu(139700)

    add_run("The feedback was collected from the students during ")
    add_run(date_str, bold=True, italic=True)
    add_run(f" (for {semester}), post completion of evaluation and assessment.")
    doc.add_paragraph()

    # --- Section 3: ATR table ---
    num_qs = len(question_data)
    table  = doc.add_table(rows=num_qs + 1, cols=3)
    _set_table_width_centered(table, 10800)
    _set_table_borders(table, val="single", sz="4", color="000000")

    COL_WIDTHS = [565, 4094, 6141]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            _set_col_width(cell, COL_WIDTHS[i])

    def set_atr_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = align
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.bold = bold
        _set_cell_valign(cell, "center")

    hdr = table.rows[0]
    set_atr_cell(hdr.cells[0], "Sl. No.",      bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_atr_cell(hdr.cells[1], "Feedback",     bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_atr_cell(hdr.cells[2], "Action Taken", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for i, (q, action_text) in enumerate(zip(question_data, action_taken_texts)):
        row = table.rows[i + 1]

        sl = row.cells[0]
        sl.text = ""
        sl_p = sl.paragraphs[0]
        sl_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sl_p.add_run(str(i + 1)).font.size = Pt(11)
        _set_cell_valign(sl, "center")

        fb = row.cells[1]
        fb.text = ""
        fb_p = fb.paragraphs[0]
        fb_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fb_p.add_run(q["description"]).font.size = Pt(11)
        _set_cell_valign(fb, "center")

        at = row.cells[2]
        at.text = ""
        at_p = at.paragraphs[0]
        at_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        at_p.add_run(action_text).font.size = Pt(11)
        _set_cell_valign(at, "center")

    # --- Section 4: Signature ---
    doc.add_paragraph()
    doc.add_paragraph()
    sig1 = doc.add_paragraph()
    sig1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig1.add_run("     \u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026..     ").font.size = Pt(11)
    sig2 = doc.add_paragraph()
    sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig2.add_run(
        "(Ratified by PAC \u2013 Prof. Nairanjana Chowdhury,  "
        "Prof. Mrinal Kanti Nath & Subhadra Shaw)"
    ).font.size = Pt(11)

    return _save_and_inject(doc, output_path, "ATR document")
